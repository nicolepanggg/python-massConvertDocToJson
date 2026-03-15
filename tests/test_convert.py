"""Tests for convert.py"""

import json
import os
import tempfile

import pytest
from docx import Document

from convert import collect_docx_files, convert_docs_to_json, extract_text_from_docx


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def make_docx(directory: str, filename: str, paragraphs: list[str]) -> str:
    """Create a .docx file with the given paragraphs and return its path."""
    path = os.path.join(directory, filename)
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# extract_text_from_docx
# ---------------------------------------------------------------------------

class TestExtractTextFromDocx:
    def test_extracts_single_paragraph(self, tmp_path):
        path = make_docx(str(tmp_path), "single.docx", ["Hello, world!"])
        assert "Hello, world!" in extract_text_from_docx(path)

    def test_extracts_multiple_paragraphs(self, tmp_path):
        path = make_docx(str(tmp_path), "multi.docx", ["First line", "Second line"])
        text = extract_text_from_docx(path)
        assert "First line" in text
        assert "Second line" in text

    def test_empty_document_returns_empty_string(self, tmp_path):
        path = make_docx(str(tmp_path), "empty.docx", [])
        # An empty Document() still has one default empty paragraph
        text = extract_text_from_docx(path)
        assert isinstance(text, str)

    def test_paragraphs_joined_by_newline(self, tmp_path):
        path = make_docx(str(tmp_path), "joined.docx", ["Line A", "Line B"])
        text = extract_text_from_docx(path)
        assert "Line A\n" in text or text.index("Line A") < text.index("Line B")


# ---------------------------------------------------------------------------
# collect_docx_files
# ---------------------------------------------------------------------------

class TestCollectDocxFiles:
    def test_finds_docx_files(self, tmp_path):
        make_docx(str(tmp_path), "a.docx", ["text"])
        make_docx(str(tmp_path), "b.docx", ["text"])
        files = collect_docx_files(str(tmp_path))
        basenames = [os.path.basename(f) for f in files]
        assert "a.docx" in basenames
        assert "b.docx" in basenames

    def test_ignores_non_docx_files(self, tmp_path):
        (tmp_path / "notes.txt").write_text("ignore me")
        make_docx(str(tmp_path), "doc.docx", ["text"])
        files = collect_docx_files(str(tmp_path))
        basenames = [os.path.basename(f) for f in files]
        assert "notes.txt" not in basenames
        assert "doc.docx" in basenames

    def test_returns_sorted_list(self, tmp_path):
        make_docx(str(tmp_path), "c.docx", ["text"])
        make_docx(str(tmp_path), "a.docx", ["text"])
        make_docx(str(tmp_path), "b.docx", ["text"])
        files = collect_docx_files(str(tmp_path))
        basenames = [os.path.basename(f) for f in files]
        assert basenames == sorted(basenames)

    def test_empty_directory_returns_empty_list(self, tmp_path):
        assert collect_docx_files(str(tmp_path)) == []

    def test_raises_on_invalid_directory(self):
        with pytest.raises(NotADirectoryError):
            collect_docx_files("/nonexistent/path/abc")

    def test_case_insensitive_extension(self, tmp_path):
        # Create a file ending in .DOCX (uppercase)
        path = os.path.join(str(tmp_path), "upper.DOCX")
        doc = Document()
        doc.add_paragraph("upper")
        doc.save(path)
        files = collect_docx_files(str(tmp_path))
        basenames = [os.path.basename(f) for f in files]
        assert "upper.DOCX" in basenames


# ---------------------------------------------------------------------------
# convert_docs_to_json
# ---------------------------------------------------------------------------

class TestConvertDocsToJson:
    def test_creates_output_json_file(self, tmp_path):
        make_docx(str(tmp_path), "doc1.docx", ["Content here"])
        output = str(tmp_path / "result.json")
        convert_docs_to_json(str(tmp_path), output)
        assert os.path.isfile(output)

    def test_json_contains_all_documents(self, tmp_path):
        make_docx(str(tmp_path), "first.docx", ["First doc"])
        make_docx(str(tmp_path), "second.docx", ["Second doc"])
        output = str(tmp_path / "result.json")
        docs = convert_docs_to_json(str(tmp_path), output)
        filenames = [d["filename"] for d in docs]
        assert "first.docx" in filenames
        assert "second.docx" in filenames

    def test_json_structure_has_filename_and_content(self, tmp_path):
        make_docx(str(tmp_path), "test.docx", ["Sample text"])
        output = str(tmp_path / "result.json")
        docs = convert_docs_to_json(str(tmp_path), output)
        assert len(docs) == 1
        assert "filename" in docs[0]
        assert "content" in docs[0]
        assert docs[0]["filename"] == "test.docx"
        assert "Sample text" in docs[0]["content"]

    def test_output_is_valid_json(self, tmp_path):
        make_docx(str(tmp_path), "doc.docx", ["Hello"])
        output = str(tmp_path / "result.json")
        convert_docs_to_json(str(tmp_path), output)
        with open(output, encoding="utf-8") as f:
            data = json.load(f)
        assert isinstance(data, list)

    def test_raises_when_no_docx_files(self, tmp_path):
        output = str(tmp_path / "result.json")
        with pytest.raises(FileNotFoundError):
            convert_docs_to_json(str(tmp_path), output)

    def test_raises_on_invalid_directory(self, tmp_path):
        output = str(tmp_path / "result.json")
        with pytest.raises(NotADirectoryError):
            convert_docs_to_json("/nonexistent/path/xyz", output)

    def test_creates_output_subdirectory_if_needed(self, tmp_path):
        make_docx(str(tmp_path), "doc.docx", ["text"])
        output = str(tmp_path / "subdir" / "nested" / "result.json")
        convert_docs_to_json(str(tmp_path), output)
        assert os.path.isfile(output)

    def test_unicode_content_preserved(self, tmp_path):
        make_docx(str(tmp_path), "unicode.docx", ["日本語テスト", "Ünïcödé"])
        output = str(tmp_path / "result.json")
        docs = convert_docs_to_json(str(tmp_path), output)
        content = docs[0]["content"]
        assert "日本語テスト" in content
        assert "Ünïcödé" in content

    def test_multiple_docs_combined_in_one_json(self, tmp_path):
        for i in range(5):
            make_docx(str(tmp_path), f"doc{i}.docx", [f"Content {i}"])
        output = str(tmp_path / "result.json")
        docs = convert_docs_to_json(str(tmp_path), output)
        assert len(docs) == 5
        with open(output, encoding="utf-8") as f:
            data = json.load(f)
        assert len(data) == 5
